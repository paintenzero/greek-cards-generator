import yaml
import json
from pathlib import Path
from docx import Document
from docx.table import _Cell
from docx.text.run import Font, Run
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PER_ROW_VERBS = 2
ROWS_PER_PAGE = 3

def load_yaml_files(verbs_dir):
    """
    Load and merge all YAML files from the verbs directory

    :param verbs_dir: Path to the directory containing verb YAML files
    :return: List of verb dictionaries
    """
    verbs = []
    verbs_path = Path(verbs_dir)

    # Iterate through all YAML files in the verbs directory
    for yaml_file in sorted(verbs_path.glob("*.yaml")):
        with open(yaml_file, "r", encoding="utf-8") as f:
            verb_data = yaml.safe_load(f)
            if isinstance(verb_data, list):
                verbs.extend(verb_data)
            else:
                verbs.append(verb_data)

    return verbs

def set_table_header_bg_color(cell):
    """
    set Background shading for Cell
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "1C5D99")
    tblCellProperties.append(clShading)
    return cell

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_margins(cell: _Cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end", ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

verbs = load_yaml_files("verbs")

tenses = [
    'present_tense',
    'simple_future_tense',
    'simple_past_tense',
    'imperative_simple_singular',
    'imperative_simple_plural',
    'imperative_negation_singular',
    'imperative_negation_plural'
]
languages = [ "english", "russian" ]

for lang in languages:

    document = Document()
    # Set style for word card header
    style = document.styles['Subtitle']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(14)
    font.color.rgb = RGBColor(255,255,255)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    if lang == "russian":
        document.add_heading('Греческие глаголы A1', 0)
    else:
        document.add_heading('Greek Verbs A1', 0)

    # Create a table for a word
    for i in range(0, len(verbs), PER_ROW_VERBS):
        if i % ROWS_PER_PAGE * PER_ROW_VERBS == 0:
            if i != 0:
                document.add_page_break()
            current_row_table = document.add_table(rows=0, cols=PER_ROW_VERBS)

        current_row = current_row_table.add_row().cells

        for j in range(0, PER_ROW_VERBS):
            set_cell_margins(current_row[j], start=0)
            if len(verbs) <= i+j:
                continue
            verb = verbs[i+j]
            delete_paragraph(current_row[j].paragraphs[0])
            table = current_row[j].add_table(rows=1, cols=2)
            table.rows[0].height = Pt(20)
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            hdr_cells = table.rows[0].cells
            hdr_cells[0].merge(hdr_cells[1])
            p = hdr_cells[0].paragraphs[0]
            p.text = verb['verb']
            p.style = 'Subtitle'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_table_header_bg_color(hdr_cells[0])

            for tense in tenses:
                if tense not in verb['tenses']:
                    continue
                row = table.add_row()
                row.height = Pt(14)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row_cells = row.cells
                set_cell_border(row_cells[0], start={"val": "single", "color": "ffffff"}, end={"val": "single", "color": "ffffff"})
                set_cell_border(row_cells[1], start={"val": "single", "color": "ffffff"}, end={"val": "single", "color": "ffffff"})
                row_cells[0].text = verb['tenses'][tense][0]['greek']
                row_cells[1].text = verb['tenses'][tense][0][lang]

    document.save(f'verbs_{lang}.docx')
